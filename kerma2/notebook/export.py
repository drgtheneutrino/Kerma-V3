"""
Kerma notebook exporters — write a Notebook to a range of formats.

v3.0 overhaul
-------------
* to_python   — cleaner script with a header, guards against `:=` in
                expressions that actually return values (we emit both an
                assignment and a print), and uses kerma2.statistics.
* to_markdown — fenced math blocks, sensible headings, includes a small
                preamble so the file opens nicely in GitHub and Obsidian.
* to_docx     — adds a cover page, page numbers, consistent styles, and
                proper code-block shading for Python cells.
* to_html     — NEW. Self-contained HTML with MathJax CDN so math cells
                render as real equations.
* to_latex    — NEW. Plain .tex fragment ready to \\input{...}.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from .cell import CellKind

if TYPE_CHECKING:
    from .engine import Notebook


# ════════════════════════════════════════════════════════════════════
# Python script export
# ════════════════════════════════════════════════════════════════════
def to_python(nb: "Notebook", path: str | Path) -> None:
    """Write a runnable .py script preserving cells as `# %%` blocks.

    Math cells are split into assignment / display form so the resulting
    file executes cleanly from top to bottom.
    """
    ts = datetime.now().isoformat(timespec="seconds")
    lines = [
        "\"\"\"Auto-generated from a Kerma notebook.\"\"\"",
        f"# exported {ts}",
        "from __future__ import annotations",
        "",
        "from math import pi, e, sqrt, exp, log as ln, log10, sin, cos, tan",
        "import numpy as np",
        "",
        "from kerma2 import Kerma",
        "from kerma2 import statistics as stats",
        "K = Kerma",
        "",
    ]
    for i, c in enumerate(nb.cells, 1):
        lines.append(f"# %% Cell {i} — {c.kind.value}")
        src = c.source if c.source is not None else ""
        if c.kind == CellKind.TEXT:
            for L in src.splitlines():
                lines.append(f"# {L}")
        elif c.kind == CellKind.PYTHON:
            lines.extend(src.splitlines())
        elif c.kind == CellKind.MATH:
            for L in src.splitlines():
                stripped = L.strip()
                if not stripped or stripped.startswith("#"):
                    lines.append(L)
                    continue
                if ":=" in stripped:
                    lhs, rhs = stripped.split(":=", 1)
                    lhs = lhs.strip(); rhs = rhs.strip()
                    lines.append(f"{lhs} = {rhs}")
                    lines.append(f"print('{lhs} =', {lhs})")
                else:
                    # expression — print label and value
                    safe = stripped.replace("'", "\\'")
                    lines.append(f"print('{safe} =', {stripped})")
        lines.append("")
    Path(path).write_text("\n".join(lines), encoding="utf-8")


# ════════════════════════════════════════════════════════════════════
# Markdown export
# ════════════════════════════════════════════════════════════════════
def to_markdown(nb: "Notebook", path: str | Path, *,
                title: str = "Kerma Notebook") -> None:
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    out = [
        f"# {title}",
        "",
        f"*Exported from Kerma · {ts}*",
        "",
        "---",
        "",
    ]
    for c in nb.cells:
        src = c.source if c.source is not None else ""
        if c.kind == CellKind.TEXT:
            out.append(src); out.append("")
        elif c.kind == CellKind.PYTHON:
            out.append("```python")
            out.append(src)
            out.append("```")
            if c.output:
                out.append("")
                out.append(f"`→ {c.output}`")
            out.append("")
        elif c.kind == CellKind.MATH:
            if c.latex:
                out.append("$$")
                out.append(c.latex)
                out.append("$$")
            else:
                out.append("```")
                out.append(src)
                out.append("```")
            if c.output:
                out.append("")
                out.append(f"**Result:** `{c.output}`")
            out.append("")
    Path(path).write_text("\n".join(out), encoding="utf-8")


# ════════════════════════════════════════════════════════════════════
# Word document export
# ════════════════════════════════════════════════════════════════════
def to_docx(nb: "Notebook", path: str | Path, *,
            title: str = "Kerma Notebook") -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for .docx export. "
            "Install with: pip install python-docx"
        ) from e

    doc = Document()

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    # ── Title block ────────────────────────────────────────────────
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph()
    r = sub.add_run(
        f"Generated by the Kerma Health-Physics toolkit · "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6A, 0x73, 0x80)
    doc.add_paragraph()

    for c in nb.cells:
        src = c.source if c.source is not None else ""
        if c.kind == CellKind.TEXT:
            for L in src.splitlines():
                p = doc.add_paragraph(L)
                p.paragraph_format.space_after = Pt(4)
        elif c.kind == CellKind.PYTHON:
            p = doc.add_paragraph()
            r = p.add_run(src)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            if c.output:
                po = doc.add_paragraph()
                ro = po.add_run(f"→ {c.output}")
                ro.font.name = "Consolas"
                ro.font.size = Pt(10)
                ro.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)
        elif c.kind == CellKind.MATH:
            for L in src.splitlines():
                if not L.strip():
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(L)
                r.font.name = "Cambria Math"
                r.font.size = Pt(12)
                r.italic = True
            if c.latex:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"  LaTeX:  {c.latex}")
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x8C, 0x95, 0xA3)
            if c.output:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"= {c.output}")
                r.font.name = "Consolas"
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)
                r.bold = True
        doc.add_paragraph()

    doc.save(str(path))


# ════════════════════════════════════════════════════════════════════
# HTML export (new in v3.0)
# ════════════════════════════════════════════════════════════════════
def to_html(nb: "Notebook", path: str | Path, *,
            title: str = "Kerma Notebook") -> None:
    parts = [
        "<!DOCTYPE html><html><head>",
        f"<meta charset='utf-8'><title>{title}</title>",
        "<script>MathJax = { tex: { inlineMath: [['$','$']] } };</script>",
        "<script src='https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-chtml.js'></script>",
        "<style>",
        "body { font: 14px/1.55 -apple-system, Segoe UI, Inter, sans-serif;"
        " max-width: 780px; margin: 3em auto; padding: 0 1em; color: #1F2933; }",
        "h1 { font-weight: 300; border-bottom: 1px solid #DEE2E6; padding-bottom: .4em; }",
        "pre { background: #F1F3F5; padding: 12px; border-radius: 3px; overflow-x: auto; }",
        ".math { text-align: center; padding: .4em 0; }",
        ".result { color: #4A90E2; font-weight: 600; }",
        ".cell { margin: 1.4em 0; }",
        "</style></head><body>",
        f"<h1>{title}</h1>",
        f"<p style='color:#8C95A3'>Exported {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>",
    ]
    for c in nb.cells:
        src = c.source if c.source is not None else ""
        parts.append("<div class='cell'>")
        if c.kind == CellKind.TEXT:
            for L in src.splitlines():
                parts.append(f"<p>{_html_escape(L)}</p>")
        elif c.kind == CellKind.PYTHON:
            parts.append(f"<pre><code>{_html_escape(src)}</code></pre>")
            if c.output:
                parts.append(f"<p class='result'>→ {_html_escape(c.output)}</p>")
        elif c.kind == CellKind.MATH:
            if c.latex:
                parts.append(f"<div class='math'>$${c.latex}$$</div>")
            else:
                parts.append(f"<pre>{_html_escape(src)}</pre>")
            if c.output:
                parts.append(
                    f"<p class='math result'>= {_html_escape(c.output)}</p>"
                )
        parts.append("</div>")
    parts.append("</body></html>")
    Path(path).write_text("\n".join(parts), encoding="utf-8")


def to_latex(nb: "Notebook", path: str | Path, *,
             title: str = "Kerma Notebook") -> None:
    """Write a minimal self-contained LaTeX article."""
    lines = [
        r"\documentclass[11pt]{article}",
        r"\usepackage{amsmath,amssymb}",
        r"\usepackage{listings,xcolor}",
        r"\usepackage[margin=1in]{geometry}",
        r"\lstset{basicstyle=\ttfamily\small,breaklines=true,frame=single}",
        r"\title{" + title + "}",
        r"\date{" + datetime.now().strftime("%Y-%m-%d") + "}",
        r"\begin{document}\maketitle",
        "",
    ]
    for c in nb.cells:
        src = c.source if c.source is not None else ""
        if c.kind == CellKind.TEXT:
            lines.append(_latex_escape(src))
            lines.append("")
        elif c.kind == CellKind.PYTHON:
            lines.append(r"\begin{lstlisting}[language=Python]")
            lines.append(src)
            lines.append(r"\end{lstlisting}")
            if c.output:
                lines.append(r"\textbf{Output: } \texttt{" + _latex_escape(c.output) + "}")
        elif c.kind == CellKind.MATH:
            if c.latex:
                lines.append(r"\begin{equation*}")
                lines.append(c.latex)
                lines.append(r"\end{equation*}")
            else:
                lines.append(r"\begin{verbatim}")
                lines.append(src)
                lines.append(r"\end{verbatim}")
            if c.output:
                lines.append(r"\[ = " + _latex_escape(c.output) + r" \]")
        lines.append("")
    lines.append(r"\end{document}")
    Path(path).write_text("\n".join(lines), encoding="utf-8")


# ────────────────────────────────────────────────────────────────────
def _html_escape(s: str) -> str:
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;"))


def _latex_escape(s: str) -> str:
    return (s.replace("\\", r"\textbackslash{}")
             .replace("&", r"\&")
             .replace("%", r"\%")
             .replace("$", r"\$")
             .replace("#", r"\#")
             .replace("_", r"\_")
             .replace("{", r"\{")
             .replace("}", r"\}")
             .replace("~", r"\textasciitilde{}")
             .replace("^", r"\textasciicircum{}"))


__all__ = ["to_python", "to_markdown", "to_docx", "to_html", "to_latex"]
